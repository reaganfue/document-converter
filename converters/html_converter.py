"""
converters/html_converter.py — HTML 來源轉換器

支援：HTML → PDF / DOCX / MD / TXT
"""

import html as html_module
import logging
from pathlib import Path

from .base import BaseConverter
from .exceptions import ConversionError, UnsupportedFormatError
from .utils import ensure_output_dir

logger = logging.getLogger(__name__)


class HTMLConverter(BaseConverter):
    """HTML 格式來源轉換器。

    支援目標格式：pdf / docx / md / txt
    """

    source_format = "html"
    supported_targets = ["pdf", "docx", "md", "txt"]

    def convert(self, input_path: Path, output_path: Path, target_format: str) -> None:
        """執行 HTML 轉換。

        Args:
            input_path: 來源 HTML 路徑
            output_path: 目標檔案路徑
            target_format: "pdf" | "docx" | "md" | "txt"

        Raises:
            UnsupportedFormatError: target_format 不在支援清單
            ConversionError: 轉換失敗
        """
        if not self.supports(target_format):
            raise UnsupportedFormatError(
                f"HTMLConverter 不支援目標格式: {target_format}"
            )

        ensure_output_dir(output_path)
        self._log_primary_attempt(target_format)

        dispatch_map = {
            "pdf": self._to_pdf,
            "docx": self._to_docx,
            "md": self._to_md,
            "txt": self._to_txt,
        }
        dispatch_map[target_format](input_path, output_path)

    # ─── 各目標格式實作 ────────────────────────────────────────────

    def _to_pdf(self, input_path: Path, output_path: Path) -> None:
        """HTML → PDF：多層渲染降級鏈（Chrome/Edge → wkhtmltopdf → weasyprint → reportlab）。"""
        try:
            from .pdf_renderer import get_renderer

            renderer = get_renderer()
            renderer.render_html_to_pdf(input_path, output_path)
            logger.info("HTML→PDF 成功")
        except Exception as exc:
            raise ConversionError("HTML→PDF 轉換失敗", source_error=exc) from exc

    def _to_docx(self, input_path: Path, output_path: Path) -> None:
        """HTML → DOCX：BeautifulSoup 解析 → python-docx 重建。"""
        try:
            html_content = input_path.read_text(encoding="utf-8", errors="replace")
            self._html_to_docx(html_content, output_path, base_dir=input_path.parent)
            logger.info("HTML→DOCX 成功")
        except Exception as exc:
            raise ConversionError("HTML→DOCX 轉換失敗", source_error=exc) from exc

    def _to_md(self, input_path: Path, output_path: Path) -> None:
        """HTML → MD：html2text。"""
        try:
            import html2text

            html_content = input_path.read_text(encoding="utf-8", errors="replace")
            converter = html2text.HTML2Text()
            converter.ignore_links = False
            converter.body_width = 0
            md_content = converter.handle(html_content)
            output_path.write_text(md_content, encoding="utf-8")
            logger.info("HTML→MD 成功")
        except Exception as exc:
            raise ConversionError("HTML→MD 轉換失敗", source_error=exc) from exc

    def _to_txt(self, input_path: Path, output_path: Path) -> None:
        """HTML → TXT：BeautifulSoup get_text()。"""
        try:
            from bs4 import BeautifulSoup

            html_content = input_path.read_text(encoding="utf-8", errors="replace")
            soup = BeautifulSoup(html_content, "html.parser")
            text = soup.get_text(separator="\n")
            # 清理多餘空行
            import re
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            output_path.write_text(text, encoding="utf-8")
            logger.info("HTML→TXT 成功")
        except Exception as exc:
            raise ConversionError("HTML→TXT 轉換失敗", source_error=exc) from exc

    # ─── HTML → DOCX 轉換邏輯 ─────────────────────────────────────

    def _html_to_docx(
        self, html_content: str, output_path: Path, base_dir: Path | None = None
    ) -> None:
        """BeautifulSoup 解析 HTML，python-docx 重建文件結構。

        支援：h1-h6 / p / ul / ol / li / table / a / strong / em / br
        """
        from bs4 import BeautifulSoup, NavigableString, Tag
        from docx import Document
        from docx.shared import Pt, RGBColor

        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        # 取 body（若無則用整個文件）
        body = soup.find("body") or soup

        for element in body.children:
            if isinstance(element, NavigableString):
                text = element.strip()
                if text:
                    doc.add_paragraph(text)
                continue

            if not isinstance(element, Tag):
                continue

            tag_name = element.name.lower()

            if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(tag_name[1])
                doc.add_heading(element.get_text(strip=True), level=level)

            elif tag_name == "p":
                para = doc.add_paragraph()
                self._add_runs_to_paragraph(para, element)

            elif tag_name in ("ul", "ol"):
                is_ordered = tag_name == "ol"
                for idx, li in enumerate(element.find_all("li", recursive=False), start=1):
                    li_text = li.get_text(strip=True)
                    prefix = f"{idx}. " if is_ordered else "• "
                    doc.add_paragraph(f"{prefix}{li_text}")

            elif tag_name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue
                cols = max(len(row.find_all(["td", "th"])) for row in rows)
                if cols == 0:
                    continue
                table = doc.add_table(rows=len(rows), cols=cols)
                for row_idx, row_el in enumerate(rows):
                    cells = row_el.find_all(["td", "th"])
                    for col_idx, cell_el in enumerate(cells):
                        if col_idx < cols:
                            table.cell(row_idx, col_idx).text = cell_el.get_text(
                                strip=True
                            )

            elif tag_name == "br":
                doc.add_paragraph()

            elif tag_name in ("pre", "code"):
                doc.add_paragraph(element.get_text())

        doc.save(str(output_path))

    def _add_runs_to_paragraph(self, para, element) -> None:
        """將 HTML 元素中的文字節點加入段落，保留 bold/italic。"""
        from bs4 import NavigableString, Tag

        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    para.add_run(text)
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                text = child.get_text()
                run = para.add_run(text)
                if tag_name in ("strong", "b"):
                    run.bold = True
                elif tag_name in ("em", "i"):
                    run.italic = True
                elif tag_name == "u":
                    run.underline = True
