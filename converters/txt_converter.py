"""
converters/txt_converter.py — 純文字來源轉換器

支援：TXT → PDF / DOCX / HTML / MD
"""

import html as html_module
import logging
import re
from pathlib import Path

from .base import BaseConverter
from .exceptions import ConversionError, UnsupportedFormatError
from .utils import build_minimal_html, ensure_output_dir, temp_file

logger = logging.getLogger(__name__)


class TXTConverter(BaseConverter):
    """純文字格式來源轉換器。

    支援目標格式：pdf / docx / html / md
    """

    source_format = "txt"
    supported_targets = ["pdf", "docx", "html", "md"]

    def convert(self, input_path: Path, output_path: Path, target_format: str) -> None:
        """執行 TXT 轉換。

        Args:
            input_path: 來源 TXT 路徑
            output_path: 目標檔案路徑
            target_format: "pdf" | "docx" | "html" | "md"

        Raises:
            UnsupportedFormatError: target_format 不在支援清單
            ConversionError: 轉換失敗
        """
        if not self.supports(target_format):
            raise UnsupportedFormatError(
                f"TXTConverter 不支援目標格式: {target_format}"
            )

        ensure_output_dir(output_path)
        self._log_primary_attempt(target_format)

        dispatch_map = {
            "pdf": self._to_pdf,
            "docx": self._to_docx,
            "html": self._to_html,
            "md": self._to_md,
        }
        dispatch_map[target_format](input_path, output_path)

    # ─── 各目標格式實作 ────────────────────────────────────────────

    def _to_pdf(self, input_path: Path, output_path: Path) -> None:
        """TXT → PDF：包裝成 <pre> HTML → 多層渲染降級鏈。"""
        try:
            from .pdf_renderer import get_renderer

            text = input_path.read_text(encoding="utf-8", errors="replace")
            escaped = html_module.escape(text)
            html_content = build_minimal_html(
                f"<pre>{escaped}</pre>",
                title=input_path.stem,
                extra_style="pre { white-space: pre-wrap; word-break: break-word; }",
            )

            with temp_file(suffix=".html") as html_path:
                html_path.write_text(html_content, encoding="utf-8")
                get_renderer().render_html_to_pdf(html_path, output_path)

            logger.info("TXT→PDF 成功")
        except Exception as exc:
            raise ConversionError("TXT→PDF 轉換失敗", source_error=exc) from exc

    def _to_docx(self, input_path: Path, output_path: Path) -> None:
        """TXT → DOCX：python-docx 逐行 add_paragraph。"""
        try:
            from docx import Document
            from docx.shared import Pt

            text = input_path.read_text(encoding="utf-8", errors="replace")
            doc = Document()

            for line in text.splitlines():
                para = doc.add_paragraph(line)
                para.style.font.size = Pt(11)

            doc.save(str(output_path))
            logger.info("TXT→DOCX 成功")
        except Exception as exc:
            raise ConversionError("TXT→DOCX 轉換失敗", source_error=exc) from exc

    def _to_html(self, input_path: Path, output_path: Path) -> None:
        """TXT → HTML：包裝 <pre> 全文。"""
        try:
            text = input_path.read_text(encoding="utf-8", errors="replace")
            escaped = html_module.escape(text)
            html_content = build_minimal_html(
                f"<pre>{escaped}</pre>",
                title=input_path.stem,
                extra_style="pre { white-space: pre-wrap; word-break: break-word; }",
            )
            output_path.write_text(html_content, encoding="utf-8")
            logger.info("TXT→HTML 成功")
        except Exception as exc:
            raise ConversionError("TXT→HTML 轉換失敗", source_error=exc) from exc

    def _to_md(self, input_path: Path, output_path: Path) -> None:
        """TXT → MD：原文 + 段落間插空行。

        連續非空行視為同一段落，空行分隔段落。
        """
        try:
            text = input_path.read_text(encoding="utf-8", errors="replace")
            # 將多個連續空行合為兩個空行（Markdown 段落分隔標準）
            md_text = re.sub(r"\n{3,}", "\n\n", text).strip()
            output_path.write_text(md_text, encoding="utf-8")
            logger.info("TXT→MD 成功")
        except Exception as exc:
            raise ConversionError("TXT→MD 轉換失敗", source_error=exc) from exc
