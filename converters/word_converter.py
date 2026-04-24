"""
converters/word_converter.py — DOCX 來源轉換器

支援：DOCX → PDF / HTML / MD / TXT
COM 安全：docx2pdf 走獨立 Lock + ThreadPoolExecutor 超時控制
降級路徑：python-docx → HTML → weasyprint
"""

import html as html_module
import logging
import threading
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from pathlib import Path

from .base import BaseConverter
from .exceptions import ConversionError, UnsupportedFormatError
from .utils import build_minimal_html, ensure_output_dir

logger = logging.getLogger(__name__)

# COM 全域鎖：序列化所有 docx2pdf 呼叫，防止 COM 競態
_DOCX2PDF_LOCK = threading.Lock()

# docx2pdf 呼叫超時（秒）
_DOCX2PDF_TIMEOUT_SEC = 60


class WordConverter(BaseConverter):
    """DOCX 格式來源轉換器。

    支援目標格式：pdf / html / md / txt
    """

    source_format = "docx"
    supported_targets = ["pdf", "html", "md", "txt"]

    def convert(self, input_path: Path, output_path: Path, target_format: str) -> None:
        """執行 DOCX 轉換。

        Args:
            input_path: 來源 DOCX 路徑
            output_path: 目標檔案路徑
            target_format: "pdf" | "html" | "md" | "txt"

        Raises:
            UnsupportedFormatError: target_format 不在支援清單
            ConversionError: 轉換失敗
        """
        if not self.supports(target_format):
            raise UnsupportedFormatError(
                f"WordConverter 不支援目標格式: {target_format}"
            )

        ensure_output_dir(output_path)
        self._log_primary_attempt(target_format)

        dispatch_map = {
            "pdf": self._to_pdf,
            "html": self._to_html,
            "md": self._to_md,
            "txt": self._to_txt,
        }
        dispatch_map[target_format](input_path, output_path)

    # ─── 各目標格式實作 ────────────────────────────────────────────

    def _to_pdf(self, input_path: Path, output_path: Path) -> None:
        """DOCX → PDF：首選 docx2pdf（COM）；降級 python-docx→HTML→weasyprint。"""
        primary_error: Exception | None = None

        # 首選：docx2pdf（COM 呼叫 + 超時保護）
        try:
            self._docx2pdf_with_timeout(input_path, output_path)
            logger.info("DOCX→PDF 首選成功: docx2pdf")
            return
        except Exception as exc:
            primary_error = exc
            self._log_fallback_attempt("pdf", exc)

        # 降級：python-docx → HTML → weasyprint
        try:
            self._docx_to_pdf_via_weasyprint(input_path, output_path)
            logger.info("DOCX→PDF 降級成功: python-docx+weasyprint")
        except Exception as fallback_error:
            self._raise_conversion_error("pdf", primary_error, fallback_error)

    def _docx2pdf_with_timeout(self, input_path: Path, output_path: Path) -> None:
        """在 ThreadPoolExecutor 單 worker 中執行 docx2pdf，含超時控制。

        COM 生命週期由 worker 函式自行管理（CoInitialize / CoUninitialize）。
        """
        def _worker() -> None:
            # 防禦性 CoInitialize：若已初始化不影響，若未初始化則補上
            try:
                import pythoncom
                pythoncom.CoInitialize()
                _initialized_here = True
            except Exception:
                _initialized_here = False

            try:
                import docx2pdf
                with _DOCX2PDF_LOCK:
                    docx2pdf.convert(str(input_path), str(output_path))
            finally:
                if _initialized_here:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_worker)
            try:
                future.result(timeout=_DOCX2PDF_TIMEOUT_SEC)
            except FutureTimeoutError:
                raise ConversionError(
                    f"docx2pdf 超時（>{_DOCX2PDF_TIMEOUT_SEC}s）"
                )
            except Exception as exc:
                raise ConversionError(f"docx2pdf 失敗: {exc}") from exc

    def _docx_to_pdf_via_weasyprint(self, input_path: Path, output_path: Path) -> None:
        """降級：python-docx → 中間 HTML → weasyprint → PDF。"""
        from .utils import temp_file

        with temp_file(suffix=".html") as html_path:
            html_content = self._docx_to_html_string(input_path)
            html_path.write_text(html_content, encoding="utf-8")
            _weasyprint_html_to_pdf(html_path, output_path)

    def _to_html(self, input_path: Path, output_path: Path) -> None:
        """DOCX → HTML：python-docx 自建 HTML 生成器。"""
        try:
            html_content = self._docx_to_html_string(input_path)
            output_path.write_text(html_content, encoding="utf-8")
            logger.info("DOCX→HTML 成功")
        except Exception as exc:
            raise ConversionError(
                "DOCX→HTML 轉換失敗",
                source_error=exc,
            ) from exc

    def _to_md(self, input_path: Path, output_path: Path) -> None:
        """DOCX → MD：DOCX → HTML → html2text。"""
        try:
            import html2text

            html_content = self._docx_to_html_string(input_path)
            converter = html2text.HTML2Text()
            converter.ignore_links = False
            converter.body_width = 0
            md_content = converter.handle(html_content)
            output_path.write_text(md_content, encoding="utf-8")
            logger.info("DOCX→MD 成功")
        except Exception as exc:
            raise ConversionError(
                "DOCX→MD 轉換失敗",
                source_error=exc,
            ) from exc

    def _to_txt(self, input_path: Path, output_path: Path) -> None:
        """DOCX → TXT：python-docx 逐段落 .text。"""
        try:
            from docx import Document

            doc = Document(str(input_path))
            lines: list[str] = []
            for para in doc.paragraphs:
                lines.append(para.text)
            output_path.write_text("\n".join(lines), encoding="utf-8")
            logger.info("DOCX→TXT 成功")
        except Exception as exc:
            raise ConversionError(
                "DOCX→TXT 轉換失敗",
                source_error=exc,
            ) from exc

    # ─── HTML 生成器（不依賴外部套件，純 python-docx） ─────────────

    def _docx_to_html_string(self, input_path: Path) -> str:
        """python-docx → HTML 字串。

        支援：標題（h1-h6）、段落、項目符號（ul/li）、有序清單（ol/li）、
              表格（table/tr/td）、粗體、斜體、超連結、換行
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document(str(input_path))
        parts: list[str] = []

        # 判斷標題層級
        _HEADING_STYLES = {
            "heading 1": "h1",
            "heading 2": "h2",
            "heading 3": "h3",
            "heading 4": "h4",
            "heading 5": "h5",
            "heading 6": "h6",
        }

        # 處理段落清單（合併連續清單項）
        in_ul = False
        in_ol = False

        def _close_list(buf: list[str]) -> None:
            nonlocal in_ul, in_ol
            if in_ul:
                buf.append("</ul>")
                in_ul = False
            if in_ol:
                buf.append("</ol>")
                in_ol = False

        def _para_to_html(para) -> str:
            """將段落轉換為 HTML 字串（不含清單標籤）。"""
            runs_html: list[str] = []
            for run in para.runs:
                text = html_module.escape(run.text)
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                if run.underline:
                    text = f"<u>{text}</u>"
                runs_html.append(text)

            # 超連結（從 XML 取 r:link）
            # python-docx 不直接暴露超連結，從 XML 取
            inner_xml = para._p.xml
            if "w:hyperlink" in inner_xml:
                # 重新處理帶有超連結的段落
                runs_html = _extract_hyperlinks(para)

            return "".join(runs_html)

        def _extract_hyperlinks(para) -> list[str]:
            """提取包含超連結的段落文字（簡化：只取文字）。"""
            parts: list[str] = []
            for run in para.runs:
                text = html_module.escape(run.text)
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                parts.append(text)
            return parts

        for para in doc.paragraphs:
            style_name = para.style.name.lower()
            para_text = _para_to_html(para)

            # 標題
            if style_name in _HEADING_STYLES:
                _close_list(parts)
                tag = _HEADING_STYLES[style_name]
                parts.append(f"<{tag}>{para_text}</{tag}>")

            # 項目符號清單
            elif style_name in ("list bullet", "list paragraph") or (
                para._p.pPr is not None and para._p.pPr.numPr is not None
            ):
                # 嘗試判斷是有序還是無序（簡化：全用 ul）
                _close_list(parts)
                if not in_ul:
                    parts.append("<ul>")
                    in_ul = True
                parts.append(f"<li>{para_text}</li>")

            # 一般段落
            else:
                _close_list(parts)
                if para_text.strip():
                    parts.append(f"<p>{para_text}</p>")
                else:
                    parts.append("<br>")

        _close_list(parts)

        # 處理表格
        table_html: list[str] = []
        for table in doc.tables:
            table_html.append("<table>")
            for row_idx, row in enumerate(table.rows):
                table_html.append("<tr>")
                for cell in row.cells:
                    tag = "th" if row_idx == 0 else "td"
                    cell_text = html_module.escape(cell.text)
                    table_html.append(f"<{tag}>{cell_text}</{tag}>")
                table_html.append("</tr>")
            table_html.append("</table>")

        body = "\n".join(parts) + "\n" + "\n".join(table_html)
        return build_minimal_html(body, title=input_path.stem)


def _weasyprint_html_to_pdf(html_path: Path, output_path: Path) -> None:
    """用 weasyprint 將 HTML 檔案轉換為 PDF。

    此函式也被 markdown_converter 和 txt_converter 呼叫。
    """
    from weasyprint import HTML as WeasyHTML

    WeasyHTML(filename=str(html_path)).write_pdf(str(output_path))
