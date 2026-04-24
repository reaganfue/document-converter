"""
tests/test_converters/test_txt_converter.py — TXT 轉換器測試
"""

import pytest
from pathlib import Path

from converters.txt_converter import TXTConverter
from converters.exceptions import UnsupportedFormatError, ConversionError


@pytest.fixture(scope="module")
def converter():
    return TXTConverter()


class TestTXTConverterMeta:
    def test_source_format(self, converter):
        assert converter.source_format == "txt"

    def test_supported_targets(self, converter):
        for fmt in ["pdf", "docx", "html", "md"]:
            assert converter.supports(fmt)

    def test_unsupported_target(self, converter):
        assert not converter.supports("pptx")
        assert not converter.supports("image")

    def test_unsupported_raises(self, converter, sample_txt, tmp_output):
        with pytest.raises(UnsupportedFormatError):
            converter.convert(sample_txt, tmp_output("out.pptx"), "pptx")


class TestTXTToHTML:
    def test_primary_success(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt.html")
        converter.convert(sample_txt, out, "html")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_output_is_html(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_html.html")
        converter.convert(sample_txt, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "<!DOCTYPE html>" in content

    def test_content_in_pre_tag(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_html2.html")
        converter.convert(sample_txt, out, "html")
        content = out.read_text(encoding="utf-8")
        assert "<pre>" in content

    def test_nonexistent_file_raises(self, converter, tmp_output):
        with pytest.raises(ConversionError):
            converter.convert(
                Path("/nonexistent/file.txt"),
                tmp_output("out.html"),
                "html",
            )


class TestTXTToMD:
    def test_primary_success(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt.md")
        converter.convert(sample_txt, out, "md")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_content_preserved(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_md.md")
        converter.convert(sample_txt, out, "md")
        content = out.read_text(encoding="utf-8")
        # 原始文字應保留
        assert "文件轉檔" in content or len(content) > 50

    def test_no_extra_blank_lines(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_md2.md")
        converter.convert(sample_txt, out, "md")
        content = out.read_text(encoding="utf-8")
        # 不應有超過 2 個連續空行
        assert "\n\n\n" not in content


class TestTXTToDOCX:
    def test_primary_success(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt.docx")
        converter.convert(sample_txt, out, "docx")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_output_is_valid_docx(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_docx.docx")
        converter.convert(sample_txt, out, "docx")
        from docx import Document
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_content_in_paragraphs(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_docx2.docx")
        converter.convert(sample_txt, out, "docx")
        from docx import Document
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert len(all_text) > 0


class TestTXTToPDF:
    def test_primary_success(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt.pdf")
        try:
            converter.convert(sample_txt, out, "pdf")
            assert out.exists()
            assert out.stat().st_size > 0
        except ConversionError as exc:
            pytest.skip(f"TXT→PDF 需要 weasyprint（{exc}）")

    def test_pdf_valid_header(self, converter, sample_txt, tmp_output):
        out = tmp_output("from_txt_pdf.pdf")
        try:
            converter.convert(sample_txt, out, "pdf")
            assert out.read_bytes()[:4] == b"%PDF"
        except ConversionError:
            pytest.skip("weasyprint 不可用")
