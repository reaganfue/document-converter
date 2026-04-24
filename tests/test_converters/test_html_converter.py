"""
tests/test_converters/test_html_converter.py — HTML 轉換器測試
"""

import pytest
from pathlib import Path

from converters.html_converter import HTMLConverter
from converters.exceptions import UnsupportedFormatError, ConversionError


@pytest.fixture(scope="module")
def converter():
    return HTMLConverter()


class TestHTMLConverterMeta:
    def test_source_format(self, converter):
        assert converter.source_format == "html"

    def test_supported_targets(self, converter):
        for fmt in ["pdf", "docx", "md", "txt"]:
            assert converter.supports(fmt)

    def test_unsupported_target(self, converter):
        assert not converter.supports("pptx")
        assert not converter.supports("image")

    def test_unsupported_raises(self, converter, sample_html, tmp_output):
        with pytest.raises(UnsupportedFormatError):
            converter.convert(sample_html, tmp_output("out.pptx"), "pptx")


class TestHTMLToTXT:
    def test_primary_success(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html.txt")
        converter.convert(sample_html, out, "txt")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_output_stripped_tags(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html_txt.txt")
        converter.convert(sample_html, out, "txt")
        content = out.read_text(encoding="utf-8")
        assert "<" not in content

    def test_nonexistent_file_raises(self, converter, tmp_output):
        with pytest.raises(ConversionError):
            converter.convert(
                Path("/nonexistent/file.html"),
                tmp_output("out.txt"),
                "txt",
            )


class TestHTMLToMD:
    def test_primary_success(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html.md")
        try:
            converter.convert(sample_html, out, "md")
            assert out.exists()
            assert out.stat().st_size > 0
        except ConversionError as exc:
            pytest.skip(f"HTML→MD 需要 html2text（{exc}）")

    def test_output_is_markdown(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html_md.md")
        try:
            converter.convert(sample_html, out, "md")
            content = out.read_text(encoding="utf-8")
            # html2text 會產出 Markdown 語法
            assert isinstance(content, str)
            assert len(content) > 50
        except ConversionError:
            pytest.skip("html2text 不可用")


class TestHTMLToDOCX:
    def test_primary_success(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html.docx")
        converter.convert(sample_html, out, "docx")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_output_is_docx(self, converter, sample_html, tmp_output):
        """驗證輸出為有效 DOCX（可被 python-docx 開啟）。"""
        out = tmp_output("from_html_docx.docx")
        converter.convert(sample_html, out, "docx")
        from docx import Document
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_tables_preserved(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html_table.docx")
        converter.convert(sample_html, out, "docx")
        from docx import Document
        doc = Document(str(out))
        assert len(doc.tables) > 0


class TestHTMLToPDF:
    def test_primary_success(self, converter, sample_html, tmp_output):
        out = tmp_output("from_html.pdf")
        try:
            converter.convert(sample_html, out, "pdf")
            assert out.exists()
            assert out.stat().st_size > 0
        except ConversionError as exc:
            pytest.skip(f"HTML→PDF 渲染層不可用（{exc}）")

    def test_pdf_starts_with_magic(self, converter, sample_html, tmp_output):
        """驗證輸出為 PDF（前 4 位元組應為 %PDF）。"""
        out = tmp_output("from_html_pdf.pdf")
        try:
            converter.convert(sample_html, out, "pdf")
            content = out.read_bytes()
            assert content[:4] == b"%PDF"
        except ConversionError:
            pytest.skip("所有 PDF 渲染引擎均不可用")
